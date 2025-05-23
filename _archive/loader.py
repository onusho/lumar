import os
import base64
import pandas as pd
import json
import xml.etree.ElementTree as ET
from docx import Document
import win32com.client as win32
import speech_recognition as sr
from moviepy.editor import VideoFileClip
from langchain_community.document_loaders import TextLoader, PyPDFLoader, UnstructuredPowerPointLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pprint import pprint

class DataLoader:
    def __init__(self, folder_path, chunk_size=100, chunk_overlap=10):
        self.folder_path = folder_path
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def process_text_file(self, file_path):
        loader = TextLoader(file_path, encoding="utf-8")
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)
        chunks = text_splitter.split_documents(documents)
        return [
            {
                "chunk_no": int(idx),
                "text": chunk.page_content,
                "path": file_path,
                "file_type": "text",
                "media_type": "text"
            }
            for idx, chunk in enumerate(chunks)
        ]

    def process_pdf_file(self, file_path):
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size, 
            chunk_overlap=self.chunk_overlap
            )
        chunks = text_splitter.split_documents(documents)
        return [
            {
                "chunk_no": int(idx),
                "text": chunk.page_content,
                "path": file_path,
                "file_type": "pdf",
                "media_type": "text"
            }
            for idx, chunk in enumerate(chunks)
        ]

    def process_pptx_file(self, file_path):
        loader = UnstructuredPowerPointLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size, 
            chunk_overlap=self.chunk_overlap
            )
        chunks = text_splitter.split_documents(documents)
        return [
            {
                "chunk_no": int(idx),
                "text": chunk.page_content,
                "path": file_path,
                "file_type": "pptx",
                "media_type": "text"
            }
            for idx, chunk in enumerate(chunks)
        ]

    def process_xlsx_file(self, file_path):
        xls = pd.ExcelFile(file_path)
        chunked_documents = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            chunked_documents.append({
                "chunk_no": 1,
                "text": df.to_dict(),
                "path": file_path,
                "file_type": "table",
                "media_type": "text"
            })
        return chunked_documents

    def process_doc_file(self, file_path):
        try:
            word = win32.Dispatch("Word.Application")
            doc = word.Documents.Open(file_path)
            text = doc.Range().Text
            doc.Close(False)
            word.Quit()
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size, 
                chunk_overlap=self.chunk_overlap
                )
            chunks = text_splitter.split_text(text)
            return [
                {
                    "chunk_no": int(idx),
                    "text": chunk,
                    "path": file_path,
                    "file_type": "doc",
                    "media_type": "text"
                }
                for idx, chunk in enumerate(chunks)
            ]
        except Exception as e:
            print(f"Error processing DOC file {file_path}: {e}")
            return None

    def process_docx_file(self, file_path):
        try:
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            text_splitter = RecursiveCharacterTextSplitter(
                separators=[" "], 
                chunk_size=self.chunk_size, 
                chunk_overlap=self.chunk_overlap
                )
            chunks = text_splitter.split_text(text)
            return [
                {
                    "chunk_no": int(idx),
                    "text": chunk,
                    "path": file_path,
                    "file_type": "docx",
                    "media_type": "text"
                }
                for idx, chunk in enumerate(chunks)
            ]
        except Exception as e:
            print(f"Error processing DOCX file {file_path}: {e}")
            return None

    def process_audio_file(self, file_path):
        recognizer = sr.Recognizer()
        try:
            with sr.AudioFile(file_path) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data)
            chunks = [text[i:i + self.chunk_size] for i in range(0, len(text), self.chunk_size)]

            return [
                {
                    "chunk_no": int(idx),
                    "text": chunk,
                    "path": file_path,
                    "file_type": "audio",
                    "media_type": "text"
                }
                for idx, chunk in enumerate(chunks)
            ]
        except Exception as e:
            print(f"Error processing audio file {file_path}: {e}")
            return None
        
    def process_image_file(self, file_path):
        try:
            with open(file_path, 'rb') as image_file:
                image_base64 = base64.b64encode(image_file.read()).decode('utf-8')
            return [{
                "path": file_path,
                "file_type": "image",
                "image": image_base64,
                "media_type": "image"
            }]
        except Exception as e:
            print(f"Error processing image file {file_path}: {e}")
            return None

    def process_csv_file(self, file_path):
        try:
            csv_data = pd.read_csv(file_path).to_string()
            return [{
                "chunk_no": 1,
                "text": csv_data,
                "path": file_path,
                "file_type": "csv",
                "media_type": "text"
            }]
        except Exception as e:
            print(f"Error processing CSV file {file_path}: {e}")
            return None
        
    def process_xml_file(self, file_path):
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            xml_data = ET.tostring(root, encoding='unicode')
            return [{
                "chunk_no": 1,
                "text": xml_data,
                "path": file_path,
                "file_type": "xml",
                "media_type": "text"
            }]
        except Exception as e:
            print(f"Error processing XML file {file_path}: {e}")
            return None
        

    def process_json_file(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.dumps(json.load(file), indent=4)
            return [{
                "chunk_no": 1,
                "text": json_data,
                "path": file_path,
                "file_type": "json",
                "media_type": "text"
            }]
        except Exception as e:
            print(f"Error processing JSON file {file_path}: {e}")
            return None
        
    def process_video_file(self, file_path):
        video_data = []
        try:
            # Extracting audio
            video = VideoFileClip(file_path)
            audio_path = file_path.replace(os.path.splitext(file_path)[1], ".wav")
            video.audio.write_audiofile(audio_path)
            audio_data = self.process_audio_file(audio_path)
            if audio_data:
                video_data.extend(audio_data)

            # Extracting frames
            frame_interval = 1  # seconds
            for i, frame in enumerate(video.iter_frames(fps=1 // frame_interval)):
                frame_base64 = base64.b64encode(frame).decode('utf-8')
                video_data.append({
                    "chunk_no": i,
                    "image": frame_base64,
                    "path": file_path,
                    "file_type": "video_frame",
                    "media_type": "image"
                })

        except Exception as e:
            print(f"Error processing video file {file_path}: {e}")

        return video_data

    def process_file(self, file_path):
        extension = os.path.splitext(file_path)[1].lower()
        if extension == '.txt':
            return self.process_text_file(file_path)
        elif extension == '.pdf':
            return self.process_pdf_file(file_path)
        elif extension == '.pptx':
            return self.process_pptx_file(file_path)
        elif extension == '.xlsx':
            return self.process_xlsx_file(file_path)
        elif extension == '.doc':
            return self.process_doc_file(file_path)
        elif extension == '.docx':
            return self.process_docx_file(file_path)
        elif extension in ['.wav', '.mp3']:
            return self.process_audio_file(file_path)
        elif extension in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
            return self.process_image_file(file_path)
        elif extension in ['.mp4', '.avi', '.mov', '.wmv']:
            return self.process_video_file(file_path)
        elif extension == '.csv':
            return self.process_csv_file(file_path)
        elif extension == '.xml':
            return self.process_xml_file(file_path)
        elif extension == '.json':
            return self.process_json_file(file_path)
        else:
            return None

    def load_data(self):
        all_data = []
        for root, dirs, files in os.walk(self.folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                processed_data = self.process_file(file_path)
                if processed_data is not None:
                    all_data.extend(processed_data)
        return all_data






# import os
# import base64
# import pandas as pd
# import json
# import xml.etree.ElementTree as ET
# from docx import Document
# import win32com.client as win32
# import speech_recognition as sr
# from moviepy.editor import VideoFileClip
# from langchain_community.document_loaders import TextLoader, PyPDFLoader, UnstructuredPowerPointLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from pprint import pprint

# class DataLoader:
#     def __init__(self, folder_path, chunk_size=1000, chunk_overlap=100):
#         self.folder_path = folder_path
#         self.chunk_size = chunk_size
#         self.chunk_overlap = chunk_overlap

#     def process_text_file(self, file_path):
#         loader = TextLoader(file_path, encoding="utf-8")
#         documents = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)
#         chunks = text_splitter.split_documents(documents)
#         return [
#             {
#                 "chunk_no": int(idx),
#                 "text": chunk.page_content,
#                 "path": file_path,
#                 "media_type": "text"
#             }
#             for idx, chunk in enumerate(chunks)
#         ]

#     def process_pdf_file(self, file_path):
#         loader = PyPDFLoader(file_path)
#         documents = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=self.chunk_size, 
#             chunk_overlap=self.chunk_overlap
#             )
#         chunks = text_splitter.split_documents(documents)
#         return [
#             {
#                 "chunk_no": int(idx),
#                 "text": chunk.page_content,
#                 "path": file_path,
#                 "media_type": "pdf"
#             }
#             for idx, chunk in enumerate(chunks)
#         ]

#     def process_pptx_file(self, file_path):
#         loader = UnstructuredPowerPointLoader(file_path)
#         documents = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=self.chunk_size, 
#             chunk_overlap=self.chunk_overlap
#             )
#         chunks = text_splitter.split_documents(documents)
#         return [
#             {
#                 "chunk_no": int(idx),
#                 "text": chunk.page_content,
#                 "path": file_path,
#                 "media_type": "pptx"
#             }
#             for idx, chunk in enumerate(chunks)
#         ]

#     def process_xlsx_file(self, file_path):
#         xls = pd.ExcelFile(file_path)
#         chunked_documents = []
#         for sheet_name in xls.sheet_names:
#             df = pd.read_excel(file_path, sheet_name=sheet_name)
#             chunked_documents.append({
#                 "chunk_no": 1,
#                 "text": df.to_dict(),
#                 "path": file_path,
#                 "media_type": "table"
#             })
#         return chunked_documents

#     def process_doc_file(self, file_path):
#         try:
#             word = win32.Dispatch("Word.Application")
#             doc = word.Documents.Open(file_path)
#             text = doc.Range().Text
#             doc.Close(False)
#             word.Quit()
#             text_splitter = RecursiveCharacterTextSplitter(
#                 chunk_size=self.chunk_size, 
#                 chunk_overlap=self.chunk_overlap
#                 )
#             chunks = text_splitter.split_text(text)
#             return [
#                 {
#                     "chunk_no": int(idx),
#                     "text": chunk,
#                     "path": file_path,
#                     "media_type": "doc"
#                 }
#                 for idx, chunk in enumerate(chunks)
#             ]
#         except Exception as e:
#             print(f"Error processing DOC file {file_path}: {e}")
#             return None

#     def process_docx_file(self, file_path):
#         try:
#             doc = Document(file_path)
#             text = '\n'.join([para.text for para in doc.paragraphs])
#             text_splitter = RecursiveCharacterTextSplitter(
#                 separators=[" "], 
#                 chunk_size=self.chunk_size, 
#                 chunk_overlap=self.chunk_overlap
#                 )
#             chunks = text_splitter.split_text(text)
#             return [
#                 {
#                     "chunk_no": int(idx),
#                     "text": chunk,
#                     "path": file_path,
#                     "media_type": "docx"
#                 }
#                 for idx, chunk in enumerate(chunks)
#             ]
#         except Exception as e:
#             print(f"Error processing DOCX file {file_path}: {e}")
#             return None

#     def process_audio_file(self, file_path):
#         recognizer = sr.Recognizer()
#         try:
#             with sr.AudioFile(file_path) as source:
#                 audio_data = recognizer.record(source)
#                 text = recognizer.recognize_google(audio_data)
#             chunks = [text[i:i + self.chunk_size] for i in range(0, len(text), self.chunk_size)]

#             return [
#                 {
#                     "chunk_no": int(idx),
#                     "text": chunk,
#                     "path": file_path,
#                     "media_type": "audio"
#                 }
#                 for idx, chunk in enumerate(chunks)
#             ]
#         except Exception as e:
#             print(f"Error processing audio file {file_path}: {e}")
#             return None
        
#     def process_image_file(self, file_path):
#         try:
#             with open(file_path, 'rb') as image_file:
#                 image_base64 = base64.b64encode(image_file.read()).decode('utf-8')
#             return [{
#                 "path": file_path,
#                 "media_type": "image",
#                 "image": image_base64
#             }]
#         except Exception as e:
#             print(f"Error processing image file {file_path}: {e}")
#             return None

#     def process_csv_file(self, file_path):
#         try:
#             csv_data = pd.read_csv(file_path).to_string()
#             return [{
#                 "chunk_no": 1,
#                 "text": csv_data,
#                 "path": file_path,
#                 "media_type": "csv"
#             }]
#         except Exception as e:
#             print(f"Error processing CSV file {file_path}: {e}")
#             return None
        
#     def process_xml_file(self, file_path):
#         try:
#             tree = ET.parse(file_path)
#             root = tree.getroot()
#             xml_data = ET.tostring(root, encoding='unicode')
#             return [{
#                 "chunk_no": 1,
#                 "text": xml_data,
#                 "path": file_path,
#                 "media_type": "xml"
#             }]
#         except Exception as e:
#             print(f"Error processing XML file {file_path}: {e}")
#             return None
        
#     def process_json_file(self, file_path):
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 json_data = json.dumps(json.load(file), indent=4)
#             return [{
#                 "chunk_no": 1,
#                 "text": json_data,
#                 "path": file_path,
#                 "media_type": "json"
#             }]
#         except Exception as e:
#             print(f"Error processing JSON file {file_path}: {e}")
#             return None
        
#     def process_file(self, file_path):
#         extension = os.path.splitext(file_path)[1].lower()
#         if extension == '.txt':
#             return self.process_text_file(file_path)
#         elif extension == '.pdf':
#             return self.process_pdf_file(file_path)
#         elif extension == '.pptx':
#             return self.process_pptx_file(file_path)
#         elif extension == '.xlsx':
#             return self.process_xlsx_file(file_path)
#         elif extension == '.doc':
#             return self.process_doc_file(file_path)
#         elif extension == '.docx':
#             return self.process_docx_file(file_path)
#         elif extension in ['.wav', '.mp3']:
#             return self.process_audio_file(file_path)
#         elif extension in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
#             return self.process_image_file(file_path)
#         elif extension in ['.mp4', '.avi', '.mov', '.wmv']:
#             return self.process_video_file(file_path)
#         elif extension == '.csv':
#             return self.process_csv_file(file_path)
#         elif extension == '.xml':
#             return self.process_xml_file(file_path)
#         elif extension == '.json':
#             return self.process_json_file(file_path)
#         else:
#             return None
        
#     def load_data(self):
#         all_data = []
#         for root, dirs, files in os.walk(self.folder_path):
#             for file in files:
#                 file_path = os.path.join(root, file)
#                 processed_data = self.process_file(file_path)
#                 if processed_data is not None:
#                     all_data.extend(processed_data)
#         return all_data

# if __name__ == "__main__":
#     folder_path = r"C:\Users\swast\OneDrive\Desktop\Desktop\files_all"
#     # folder_path = r"C:\Users\swast\OneDrive\Desktop\Desktop\audio"
#     loader = DataLoader(folder_path)
#     data = loader.load_data()
#     pprint(data)

